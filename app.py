import io
from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
# Conversion libraries
from pypdf import PdfReader
from docx import Document
from PIL import Image

# --- Setup ---
app = Flask(__name__)
CORS(app)

# --- Route to Serve the Frontend ---
@app.route('/')
def index():
    """Serves the main HTML page."""
    return render_template('index.html')

# --- Main Conversion Logic ---
@app.route('/convert', methods=['POST'])
def convert_file():
    """Handles the file upload and conversion."""
    try:
        file = request.files['file']
        from_type = request.form['fromType']
        to_format = request.form['toFormat']
        base_filename = file.filename.rsplit('.', 1)[0]
        
        output_buffer = io.BytesIO()

        # --- Routing to the correct conversion function ---
        if from_type == 'pdf' and to_format == 'docx':
            pdf_reader = PdfReader(file.stream)
            doc = Document()
            for page in pdf_reader.pages:
                doc.add_paragraph(page.extract_text() or "") # Use empty string if text is None
            doc.save(output_buffer)
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            download_name = f"{base_filename}.docx"

        elif from_type == 'image' and to_format == 'pdf':
            img = Image.open(file.stream).convert("RGB")
            img.save(output_buffer, format='PDF', resolution=100.0)
            mime_type = 'application/pdf'
            download_name = f"{base_filename}.pdf"

        elif from_type == 'text' and to_format == 'docx':
            text_content = file.read().decode('utf-8')
            doc = Document()
            doc.add_paragraph(text_content)
            doc.save(output_buffer)
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            download_name = f"{base_filename}.docx"
            
        else:
            return jsonify({"error": f"Conversion from {from_type} to {to_format} is not supported."}), 400

        output_buffer.seek(0)
        return send_file(output_buffer, as_attachment=True, download_name=download_name, mimetype=mime_type)

    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500

# This part is for local testing; Gunicorn will run the app in production
if __name__ == '__main__':
    app.run(debug=True, port=5000)
